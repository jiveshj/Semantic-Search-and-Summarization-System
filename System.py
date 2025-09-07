# Semantic Search and Summarization System
# A scalable document search and summarization platform using transformer models

import os
import json
import pickle
import logging
from typing import List, Dict, Tuple, Optional
from dataclasses import dataclass
from datetime import datetime
import asyncio
import aiohttp
from concurrent.futures import ThreadPoolExecutor
import threading

import numpy as np
import pandas as pd
import faiss
import torch
import torch.nn as nn
from torch.utils.data import DataLoader, Dataset
from transformers import (
    AutoTokenizer, AutoModel, AutoModelForSeq2SeqLM,
    BertTokenizer, BertModel, T5ForConditionalGeneration, T5Tokenizer,
    pipeline, logging as transformers_logging
)
from sentence_transformers import SentenceTransformer
import nltk
from nltk.tokenize import sent_tokenize
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.cluster import KMeans
import requests
from bs4 import BeautifulSoup
import PyPDF2
import docx
from pathlib import Path
import hashlib
import re
from tqdm import tqdm

# Suppress transformers warnings
transformers_logging.set_verbosity_error()
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class Document:
    """Document data structure"""
    id: str
    title: str
    content: str
    source: str
    metadata: Dict
    embedding: Optional[np.ndarray] = None
    created_at: datetime = None
    
    def __post_init__(self):
        if self.created_at is None:
            self.created_at = datetime.now()

@dataclass
class SearchResult:
    """Search result data structure"""
    document: Document
    score: float
    relevance_explanation: str
    summary: Optional[str] = None

class DocumentProcessor:
    """Handles document ingestion and preprocessing"""
    
    def __init__(self):
        self.supported_formats = ['.txt', '.pdf', '.docx', '.html', '.md']
        
    def extract_text_from_file(self, file_path: str) -> str:
        """Extract text from various file formats"""
        path = Path(file_path)
        extension = path.suffix.lower()
        
        try:
            if extension == '.txt' or extension == '.md':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
                    
            elif extension == '.pdf':
                return self._extract_from_pdf(file_path)
                
            elif extension == '.docx':
                return self._extract_from_docx(file_path)
                
            elif extension == '.html':
                return self._extract_from_html(file_path)
                
            else:
                logger.warning(f"Unsupported file format: {extension}")
                return ""
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return ""
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF files"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            return ""
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            doc = docx.Document(file_path)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            return ""
    
    def _extract_from_html(self, file_path: str) -> str:
        """Extract text from HTML files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
                return soup.get_text()
        except Exception as e:
            logger.error(f"HTML extraction error: {str(e)}")
            return ""
    
    def clean_text(self, text: str) -> str:
        """Clean and preprocess text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s\.,!?;:()\-]', '', text)
        return text.strip()
    
    def chunk_document(self, text: str, chunk_size: int = 512, overlap: int = 50) -> List[str]:
        """Split document into overlapping chunks for better processing"""
        sentences = sent_tokenize(text)
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) <= chunk_size:
                current_chunk += " " + sentence
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence
                
        if current_chunk:
            chunks.append(current_chunk.strip())
            
        return chunks

class EmbeddingModel:
    """Handles document embeddings using transformer models"""
    
    def __init__(self, model_name: str = 'sentence-transformers/all-MiniLM-L6-v2'):
        self.device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
        logger.info(f"Using device: {self.device}")
        
        # Load sentence transformer model
        self.model = SentenceTransformer(model_name)
        self.model.to(self.device)
        
        # Alternative: Use BERT directly for more control
        # self.tokenizer = AutoTokenizer.from_pretrained('bert-base-uncased')
        # self.bert_model = AutoModel.from_pretrained('bert-base-uncased')
        
        self.embedding_dim = self.model.get_sentence_embedding_dimension()
        logger.info(f"Embedding dimension: {self.embedding_dim}")
    
    def encode_documents(self, texts: List[str], batch_size: int = 32) -> np.ndarray:
        """Generate embeddings for a batch of documents"""
        embeddings = []
        
        for i in tqdm(range(0, len(texts), batch_size), desc="Generating embeddings"):
            batch = texts[i:i + batch_size]
            batch_embeddings = self.model.encode(
                batch,
                convert_to_numpy=True,
                show_progress_bar=False,
                batch_size=batch_size
            )
            embeddings.extend(batch_embeddings)
        
        return np.array(embeddings)
    
    def encode_query(self, query: str) -> np.ndarray:
        """Generate embedding for a single query"""
        return self.model.encode([query], convert_to_numpy=True)[0]

class VectorIndex:
    """FAISS-based vector index for efficient similarity search"""
    
    def __init__(self, embedding_dim: int, index_type: str = 'flat'):
        self.embedding_dim = embedding_dim
        self.index_type = index_type
        self.documents = []  # Store document metadata
        
        # Initialize FAISS index
        if index_type == 'flat':
            self.index = faiss.IndexFlatIP(embedding_dim)  # Inner Product
        elif index_type == 'ivf':
            # Inverted File Index for larger datasets
            quantizer = faiss.IndexFlatIP(embedding_dim)
            self.index = faiss.IndexIVFFlat(quantizer, embedding_dim, 100)
        else:
            raise ValueError(f"Unsupported index type: {index_type}")
        
        logger.info(f"Created FAISS index: {index_type}, dimension: {embedding_dim}")
    
    def add_documents(self, documents: List[Document], embeddings: np.ndarray):
        """Add documents and their embeddings to the index"""
        # Normalize embeddings for cosine similarity
        faiss.normalize_L2(embeddings)
        
        if self.index_type == 'ivf' and not self.index.is_trained:
            # Train IVF index if not already trained
            self.index.train(embeddings)
        
        self.index.add(embeddings)
        self.documents.extend(documents)
        
        logger.info(f"Added {len(documents)} documents to index. Total: {self.index.ntotal}")
    
    def search(self, query_embedding: np.ndarray, k: int = 10) -> List[Tuple[int, float]]:
        """Search for similar documents"""
        # Normalize query embedding
        query_embedding = query_embedding.reshape(1, -1)
        faiss.normalize_L2(query_embedding)
        
        scores, indices = self.index.search(query_embedding, k)
        
        results = []
        for i, (idx, score) in enumerate(zip(indices[0], scores[0])):
            if idx >= 0:  # Valid index
                results.append((idx, float(score)))
        
        return results
    
    def save_index(self, filepath: str):
        """Save the index to disk"""
        faiss.write_index(self.index, f"{filepath}.index")
        with open(f"{filepath}.docs", 'wb') as f:
            pickle.dump(self.documents, f)
        logger.info(f"Index saved to {filepath}")
    
    def load_index(self, filepath: str):
        """Load the index from disk"""
        self.index = faiss.read_index(f"{filepath}.index")
        with open(f"{filepath}.docs", 'rb') as f:
            self.documents = pickle.load(f)
        logger.info(f"Index loaded from {filepath}")

class Summarizer:
    """Document summarization using transformer models"""
    
    def __init__(self, model_name: str = 'facebook/bart-large-cnn'):
        self.device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
        
        # Load summarization model
        self.summarizer = pipeline(
            "summarization",
            model=model_name,
            device=0 if torch.cuda.is_available() else -1,
            framework="pt"
        )
        
        logger.info(f"Loaded summarization model: {model_name}")
    
    def summarize_text(self, text: str, max_length: int = 150, min_length: int = 50) -> str:
        """Generate summary for a given text"""
        try:
            # Handle long texts by chunking
            if len(text.split()) > 1000:
                text = self._truncate_text(text, 1000)
            
            summary = self.summarizer(
                text,
                max_length=max_length,
                min_length=min_length,
                do_sample=False,
                truncation=True
            )
            
            return summary[0]['summary_text']
            
        except Exception as e:
            logger.error(f"Summarization error: {str(e)}")
            return self._extractive_summary(text, max_length)
    
    def _truncate_text(self, text: str, max_words: int) -> str:
        """Truncate text to specified number of words"""
        words = text.split()
        return ' '.join(words[:max_words])
    
    def _extractive_summary(self, text: str, max_length: int) -> str:
        """Fallback extractive summarization"""
        sentences = sent_tokenize(text)
        if len(sentences) <= 3:
            return text
        
        # Simple extractive summary - take first and most central sentences
        summary_sentences = sentences[:2]
        if len(sentences) > 5:
            summary_sentences.append(sentences[len(sentences)//2])
        
        summary = ' '.join(summary_sentences)
        return summary[:max_length] + "..." if len(summary) > max_length else summary

class SemanticSearchSystem:
    """Main semantic search and summarization system"""
    
    def __init__(self, embedding_model_name: str = 'sentence-transformers/all-MiniLM-L6-v2'):
        self.processor = DocumentProcessor()
        self.embedding_model = EmbeddingModel(embedding_model_name)
        self.vector_index = VectorIndex(self.embedding_model.embedding_dim)
        self.summarizer = Summarizer()
        
        # Performance tracking
        self.query_stats = {
            'total_queries': 0,
            'average_response_time': 0.0,
            'cache_hits': 0
        }
        
        # Simple query cache
        self.query_cache = {}
        self.max_cache_size = 1000
        
        logger.info("Semantic Search System initialized")
    
    def ingest_documents(self, document_paths: List[str], batch_size: int = 100) -> int:
        """Ingest documents from file paths"""
        documents = []
        all_texts = []
        
        logger.info(f"Processing {len(document_paths)} documents...")
        
        for path in tqdm(document_paths, desc="Processing documents"):
            try:
                # Extract text
                raw_text = self.processor.extract_text_from_file(path)
                if not raw_text:
                    continue
                
                # Clean text
                cleaned_text = self.processor.clean_text(raw_text)
                
                # Create document
                doc_id = hashlib.md5(path.encode()).hexdigest()
                document = Document(
                    id=doc_id,
                    title=Path(path).stem,
                    content=cleaned_text,
                    source=path,
                    metadata={
                        'file_size': os.path.getsize(path),
                        'file_type': Path(path).suffix,
                        'word_count': len(cleaned_text.split())
                    }
                )
                
                documents.append(document)
                all_texts.append(cleaned_text)
                
            except Exception as e:
                logger.error(f"Error processing {path}: {str(e)}")
                continue
        
        if not documents:
            logger.warning("No documents were successfully processed")
            return 0
        
        # Generate embeddings
        logger.info("Generating embeddings...")
        embeddings = self.embedding_model.encode_documents(all_texts, batch_size)
        
        # Add to vector index
        self.vector_index.add_documents(documents, embeddings)
        
        logger.info(f"Successfully ingested {len(documents)} documents")
        return len(documents)
    
    def add_document_from_text(self, title: str, content: str, source: str = "manual") -> str:
        """Add a single document from text content"""
        doc_id = hashlib.md5(f"{title}{content}".encode()).hexdigest()
        
        document = Document(
            id=doc_id,
            title=title,
            content=content,
            source=source,
            metadata={'word_count': len(content.split())}
        )
        
        # Generate embedding
        embedding = self.embedding_model.encode_documents([content], batch_size=1)
        
        # Add to index
        self.vector_index.add_documents([document], embedding)
        
        return doc_id
    
    def search(self, query: str, k: int = 10, include_summary: bool = True) -> List[SearchResult]:
        """Perform semantic search"""
        start_time = datetime.now()
        
        # Check cache
        cache_key = f"{query}_{k}_{include_summary}"
        if cache_key in self.query_cache:
            self.query_stats['cache_hits'] += 1
            return self.query_cache[cache_key]
        
        try:
            # Generate query embedding
            query_embedding = self.embedding_model.encode_query(query)
            
            # Search vector index
            search_results = self.vector_index.search(query_embedding, k)
            
            # Convert to SearchResult objects
            results = []
            for doc_idx, score in search_results:
                document = self.vector_index.documents[doc_idx]
                
                # Generate relevance explanation
                explanation = self._generate_relevance_explanation(query, document.content, score)
                
                # Generate summary if requested
                summary = None
                if include_summary:
                    summary = self.summarizer.summarize_text(document.content)
                
                search_result = SearchResult(
                    document=document,
                    score=score,
                    relevance_explanation=explanation,
                    summary=summary
                )
                results.append(search_result)
            
            # Update cache
            if len(self.query_cache) < self.max_cache_size:
                self.query_cache[cache_key] = results
            
            # Update stats
            self.query_stats['total_queries'] += 1
            response_time = (datetime.now() - start_time).total_seconds()
            self.query_stats['average_response_time'] = (
                (self.query_stats['average_response_time'] * (self.query_stats['total_queries'] - 1) + response_time)
                / self.query_stats['total_queries']
            )
            
            logger.info(f"Search completed in {response_time:.3f}s, found {len(results)} results")
            return results
            
        except Exception as e:
            logger.error(f"Search error: {str(e)}")
            return []
    
    def _generate_relevance_explanation(self, query: str, content: str, score: float) -> str:
        """Generate explanation for why a document is relevant"""
        # Simple keyword matching for explanation
        query_terms = set(query.lower().split())
        content_terms = set(content.lower().split())
        
        common_terms = query_terms.intersection(content_terms)
        
        if common_terms:
            return f"Matches query terms: {', '.join(list(common_terms)[:5])} (similarity: {score:.3f})"
        else:
            return f"Semantic similarity score: {score:.3f}"
    
    def get_document_by_id(self, doc_id: str) -> Optional[Document]:
        """Retrieve document by ID"""
        for doc in self.vector_index.documents:
            if doc.id == doc_id:
                return doc
        return None
    
    def get_stats(self) -> Dict:
        """Get system statistics"""
        return {
            'total_documents': len(self.vector_index.documents),
            'index_size': self.vector_index.index.ntotal,
            'embedding_dimension': self.embedding_model.embedding_dim,
            'query_stats': self.query_stats,
            'cache_size': len(self.query_cache)
        }
    
    def save_system(self, filepath: str):
        """Save the entire system to disk"""
        self.vector_index.save_index(filepath)
        
        # Save additional system state
        system_state = {
            'query_stats': self.query_stats,
            'query_cache': self.query_cache
        }
        
        with open(f"{filepath}.state", 'wb') as f:
            pickle.dump(system_state, f)
        
        logger.info(f"System saved to {filepath}")
    
    def load_system(self, filepath: str):
        """Load the entire system from disk"""
        self.vector_index.load_index(filepath)
        
        try:
            with open(f"{filepath}.state", 'rb') as f:
                system_state = pickle.load(f)
                self.query_stats = system_state.get('query_stats', self.query_stats)
                self.query_cache = system_state.get('query_cache', {})
        except FileNotFoundError:
            logger.warning("System state file not found, using defaults")
        
        logger.info(f"System loaded from {filepath}")

# Demo and testing code
def create_sample_documents() -> List[Tuple[str, str]]:
    """Create sample documents for testing"""
    documents = [
        ("Machine Learning Basics", """
        Machine learning is a subset of artificial intelligence that focuses on algorithms 
        that can learn and make predictions from data. There are three main types: 
        supervised learning, unsupervised learning, and reinforcement learning. 
        Supervised learning uses labeled data to train models, while unsupervised learning 
        finds patterns in unlabeled data. Reinforcement learning learns through 
        interaction with an environment.
        """),
        
        ("Deep Learning Overview", """
        Deep learning is a subset of machine learning that uses neural networks with 
        multiple layers. These networks can automatically learn hierarchical 
        representations of data. Common architectures include convolutional neural 
        networks (CNNs) for image processing, recurrent neural networks (RNNs) for 
        sequential data, and transformers for natural language processing.
        """),
        
        ("Natural Language Processing", """
        Natural Language Processing (NLP) is a field that combines linguistics and 
        machine learning to help computers understand human language. Modern NLP uses 
        transformer models like BERT and GPT, which have achieved state-of-the-art 
        results in tasks like text classification, question answering, and language 
        translation. These models are pre-trained on large text corpora.
        """),
        
        ("Computer Vision Applications", """
        Computer vision enables machines to interpret visual information from the world. 
        Applications include image classification, object detection, facial recognition, 
        and autonomous vehicles. Convolutional neural networks have revolutionized this 
        field, achieving human-level performance on many visual tasks. Transfer learning 
        allows models trained on large datasets to be adapted for specific applications.
        """),
        
        ("Data Science Process", """
        Data science involves extracting insights from data through a systematic process. 
        This includes data collection, cleaning, exploratory analysis, feature engineering, 
        model building, and deployment. Tools commonly used include Python, R, SQL, and 
        various machine learning libraries. Data scientists must also consider ethical 
        implications and ensure their models are fair and interpretable.
        """),
        
        ("Cloud Computing Infrastructure", """
        Cloud computing provides on-demand access to computing resources over the internet. 
        Major providers include AWS, Azure, and Google Cloud Platform. Benefits include 
        scalability, cost-effectiveness, and global accessibility. Cloud services range 
        from basic storage and compute to advanced AI and machine learning platforms. 
        Containerization and microservices architectures are commonly used in cloud deployments.
        """),
        
        ("Software Engineering Best Practices", """
        Software engineering involves systematic approaches to designing, developing, 
        and maintaining software systems. Best practices include version control with Git, 
        automated testing, continuous integration/deployment, code reviews, and 
        documentation. Agile methodologies emphasize iterative development and 
        collaboration. Modern software architecture often uses microservices and APIs.
        """),
        
        ("Database Management Systems", """
        Database management systems (DBMS) store, organize, and manage data efficiently. 
        Relational databases like MySQL and PostgreSQL use SQL for querying structured data. 
        NoSQL databases like MongoDB handle unstructured data and scale horizontally. 
        Modern applications often use both types depending on requirements. Database design 
        principles include normalization, indexing, and transaction management.
        """),
        
        ("Cybersecurity Fundamentals", """
        Cybersecurity protects digital systems, networks, and data from threats. 
        Common threats include malware, phishing, and denial-of-service attacks. 
        Security measures include encryption, access controls, firewalls, and intrusion 
        detection systems. Security by design principles emphasize building protection 
        into systems from the beginning rather than adding it later.
        """),
        
        ("Web Development Technologies", """
        Web development involves creating applications that run in web browsers. 
        Frontend technologies include HTML, CSS, and JavaScript frameworks like React 
        and Vue.js. Backend development uses languages like Python, Node.js, and Java 
        with frameworks like Django and Express. Modern web applications are often 
        single-page applications (SPAs) that communicate with APIs.
        """)
    ]
    
    return documents

def demo_system():
    """Demonstrate the semantic search system"""
    print("🚀 Initializing Semantic Search System...")
    
    # Initialize system
    search_system = SemanticSearchSystem()
    
    # Create and add sample documents
    print("📄 Adding sample documents...")
    sample_docs = create_sample_documents()
    
    for title, content in sample_docs:
        search_system.add_document_from_text(title, content)
    
    # Demonstrate search functionality
    queries = [
        "What is machine learning?",
        "How do neural networks work?",
        "Tell me about databases",
        "What are the benefits of cloud computing?",
        "How does computer vision work?"
    ]
    
    print(f"\n🔍 Running {len(queries)} sample queries...\n")
    
    for i, query in enumerate(queries, 1):
        print(f"Query {i}: '{query}'")
        print("-" * 50)
        
        results = search_system.search(query, k=3, include_summary=True)
        
        for j, result in enumerate(results, 1):
            print(f"  Result {j}: {result.document.title}")
            print(f"  Score: {result.score:.4f}")
            print(f"  Summary: {result.summary}")
            print(f"  Explanation: {result.relevance_explanation}")
            print()
        
        print()
    
    # Display system statistics
    print("📊 System Statistics:")
    stats = search_system.get_stats()
    for key, value in stats.items():
        print(f"  {key}: {value}")

if __name__ == "__main__":
    # Download required NLTK data
    try:
        nltk.download('punkt', quiet=True)
    except:
        pass
    
    # Run demo
    demo_system()
